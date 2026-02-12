"""
Wellness Studio - Input Processing Module
Handles PDF, text, and audio transcription inputs.
"""
import io
import re
from pathlib import Path
from typing import Union, Optional, BinaryIO
import warnings

# Document processing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Audio processing
try:
    import whisper
    import librosa
    import soundfile as sf
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

from .config import processing_config, runtime_config, path_config


class InputProcessor:
    """Process various input formats (PDF, text, audio) into structured text."""

    def __init__(self):
        self.whisper_model = None
        self._supported_doc_formats = processing_config.SUPPORTED_DOC_FORMATS
        self._supported_audio_formats = processing_config.SUPPORTED_AUDIO_FORMATS

    def process_input(
        self,
        input_path: Optional[Union[str, Path]] = None,
        text_input: Optional[str] = None,
        audio_bytes: Optional[bytes] = None
    ) -> dict:
        """
        Process input from file path, raw text, or audio bytes.

        Returns:
            dict with keys: 'content', 'source_type', 'metadata'
        """
        if input_path:
            return self._process_file(input_path)
        elif text_input:
            return self._process_text(text_input)
        elif audio_bytes:
            return self._process_audio_bytes(audio_bytes)
        else:
            raise ValueError("Must provide input_path, text_input, or audio_bytes")

    def _process_file(self, file_path: Union[str, Path]) -> dict:
        """Process a file based on its extension."""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self._extract_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        elif suffix in ['.docx']:
            return self._extract_docx(file_path)
        elif suffix in self._supported_audio_formats:
            return self._transcribe_audio(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _extract_pdf(self, pdf_path: Path) -> dict:
        """Extract text from PDF using pdfplumber or PyPDF2."""
        text = ""
        metadata = {}

        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    metadata['pages'] = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {i+1} ---\n{page_text}"
            except Exception as e:
                warnings.warn(f"pdfplumber failed: {e}, trying PyPDF2")
                text = ""

        # Fallback to PyPDF2
        if not text and PYPDF2_AVAILABLE:
            try:
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata['pages'] = len(reader.pages)
                    for i, page in enumerate(reader.pages):
                        text += f"\n--- Page {i+1} ---\n{page.extract_text()}"
            except Exception as e:
                raise RuntimeError(f"Failed to extract PDF: {e}")

        if not text:
            raise RuntimeError("No PDF extraction library available")

        return {
            'content': self._clean_text(text),
            'source_type': 'pdf',
            'metadata': metadata
        }

    def _extract_text_file(self, file_path: Path) -> dict:
        """Extract text from plain text or markdown files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        return {
            'content': self._clean_text(text),
            'source_type': 'text',
            'metadata': {'lines': len(text.splitlines())}
        }

    def _extract_docx(self, file_path: Path) -> dict:
        """Extract text from Word documents."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

            return {
                'content': self._clean_text(text),
                'source_type': 'docx',
                'metadata': {'paragraphs': len(doc.paragraphs)}
            }
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

    def _process_text(self, text: str) -> dict:
        """Process raw text input."""
        return {
            'content': self._clean_text(text),
            'source_type': 'raw_text',
            'metadata': {'char_count': len(text)}
        }

    def _transcribe_audio(self, audio_path: Path) -> dict:
        """Transcribe audio file using Whisper."""
        if not WHISPER_AVAILABLE:
            raise RuntimeError(
                "Whisper not installed. Install with: pip install openai-whisper"
            )

        # Lazy load whisper model
        if self.whisper_model is None:
            self._load_whisper_model()

        result = self.whisper_model.transcribe(str(audio_path))

        return {
            'content': self._clean_text(result['text']),
            'source_type': 'audio',
            'metadata': {
                'language': result.get('language', 'unknown'),
                'duration': result.get('duration'),
                'segments': len(result.get('segments', []))
            }
        }

    def _load_whisper_model(self):
        """Load Whisper model with offline mode support."""
        offline_mode = runtime_config.LOCAL_ONLY or runtime_config.OFFLINE_MODELS
        mode_str = " (LOCAL_ONLY mode)" if offline_mode else ""
        print(f"Loading Whisper model: {processing_config.WHISPER_MODEL}{mode_str}")

        whisper_cache = path_config.MODELS_DIR / "whisper"
        whisper_cache.mkdir(parents=True, exist_ok=True)

        if offline_mode:
            # Check if model exists in cache
            model_path = whisper_cache / f"{processing_config.WHISPER_MODEL}.pt"
            if not model_path.exists():
                raise RuntimeError(
                    f"Whisper model '{processing_config.WHISPER_MODEL}' not found in cache at '{whisper_cache}'. "
                    f"In LOCAL_ONLY mode, models must be pre-downloaded. "
                    f"Run with WELLNESS_LOCAL_ONLY=false once to download, then switch back."
                )
            print(f"  Using cached model from: {whisper_cache}")

        self.whisper_model = whisper.load_model(
            processing_config.WHISPER_MODEL,
            download_root=str(whisper_cache)
        )

    def _process_audio_bytes(self, audio_bytes: bytes) -> dict:
        """Process audio from bytes (for API/integration use)."""
        if not WHISPER_AVAILABLE:
            raise RuntimeError("Whisper not installed")

        if self.whisper_model is None:
            self._load_whisper_model()

        # Write to temporary file for whisper
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name

        try:
            result = self.whisper_model.transcribe(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

        text_content = result.get('text', '')
        if isinstance(text_content, list):
            text_content = ' '.join(str(item) for item in text_content)
        return {
            'content': self._clean_text(str(text_content)),
            'source_type': 'audio_bytes',
            'metadata': {
                'language': result.get('language', 'unknown'),
                'duration': result.get('duration')
            }
        }

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text

    @property
    def supported_formats(self) -> dict:
        """Return supported input formats."""
        return {
            'document': self._supported_doc_formats,
            'audio': self._supported_audio_formats if WHISPER_AVAILABLE else ()
        }
